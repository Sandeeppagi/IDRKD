from __future__ import annotations

import shutil
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_BREAK, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SRC = Path("/Users/sandeeppagi/Downloads/IDRKD_HLD_LLD_v3.docx")
OUT = Path("/Users/sandeeppagi/Projects/IDRKD/docs/design/IDRKD_HLD_LLD_v3_updated.docx")


ACCENT = RGBColor(31, 78, 121)
MUTED = RGBColor(89, 89, 89)
LIGHT_BLUE = "D9EAF7"
LIGHT_GREY = "F2F2F2"


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    new_run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def insert_paragraph_before(anchor, text: str = "", style: str | None = None):
    p = anchor.insert_paragraph_before(text)
    if style:
        p.style = style
    return p


def make_code_style(doc: Document) -> None:
    styles = doc.styles
    if "Code Block" not in [s.name for s in styles]:
        style = styles.add_style("Code Block", 1)
    else:
        style = styles["Code Block"]
    style.font.name = "Consolas"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    style.font.size = Pt(8.5)
    style.font.color.rgb = RGBColor(34, 34, 34)
    pf = style.paragraph_format
    pf.left_indent = Inches(0.18)
    pf.right_indent = Inches(0.08)
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.0


def set_toc_update_flag(docx_path: Path) -> None:
    tmp = docx_path.with_suffix(".tmp.docx")
    with ZipFile(docx_path, "r") as zin, ZipFile(tmp, "w", ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/settings.xml":
                xml = data.decode("utf-8")
                if "<w:updateFields" not in xml:
                    xml = xml.replace(
                        "</w:settings>",
                        '<w:updateFields w:val="true"/></w:settings>',
                    )
                data = xml.encode("utf-8")
            zout.writestr(item, data)
    tmp.replace(docx_path)


def clear_footer(section) -> None:
    for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
        for p in footer.paragraphs:
            p.clear()


def rebuild_footer(section) -> None:
    clear_footer(section)
    p = section.footer.paragraphs[0]
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = p.add_run("Page ")
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fld.append(r)
    p._p.append(fld)
    run = p.add_run(" of ")
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "NUMPAGES")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fld.append(r)
    p._p.append(fld)


def style_document(doc: Document) -> None:
    make_code_style(doc)
    for style_name, size in [("Normal", 10.5), ("List Paragraph", 10.5)]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.08
    for style_name, size in [("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12)]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = ACCENT
        style.paragraph_format.keep_with_next = True
    doc.styles["Title"].font.color.rgb = RGBColor(0, 0, 0)

    code_prefixes = (
        "class ",
        "async def ",
        "def ",
        "if ",
        "return ",
        "await ",
        "MATCH ",
        "CREATE ",
        "ALTER ",
        "SELECT ",
        "@app.",
        "# ",
        "THRESHOLD",
        "CENTROID_",
        "FULL_REINDEX",
        "from ",
        "with ",
        "for ",
    )
    for p in doc.paragraphs:
        text = p.text.strip()
        if text.startswith(code_prefixes) or "BaseModel):" in text or "Literal[" in text:
            p.style = doc.styles["Code Block"]

    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)
        section.different_first_page_header_footer = False
        section.header.paragraphs[0].text = "IDRKD System - HLD & LLD"
        section.header.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        for run in section.header.paragraphs[0].runs:
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = MUTED
        rebuild_footer(section)


def add_table(rows, widths=None):
    doc = add_table.doc
    table = doc.add_table(rows=1, cols=len(rows[0]))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    hdr = table.rows[0].cells
    for i, value in enumerate(rows[0]):
        set_cell_text(hdr[i], value, bold=True)
        shade_cell(hdr[i], LIGHT_BLUE)
    for row in rows[1:]:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    return table


def append_front_matter(doc: Document) -> None:
    anchor = doc.paragraphs[7]
    # Insert in reverse because each insertion goes immediately before the anchor.
    items = []
    items.append(("break", None, None))
    items.append(("heading", "Executive Summary", "Heading 1"))
    items.append(("para", "This design proposes a six-month, single-researcher MVP for an Intelligent Data Reconciliation and Knowledge Discovery system. The core thesis claim is that MCP-native retrieval, graph reasoning, bounded critique, and SLM distillation can produce auditable, reproducible answers over enterprise technical assets without requiring a frontier model in the production path.", None))
    items.append(("heading", "MVP Scope", "Heading 2"))
    items.append(("para", "The MVP keeps all six pillars, but narrows each to the minimum needed for a defensible dissertation result: Python/JavaScript ingestion, Neo4j graph modelling, pgvector retrieval, bounded LangGraph orchestration, Phi-4-mini QLoRA distillation, and selective drift-triggered re-indexing.", None))
    items.append(("heading", "Primary Success Measures", "Heading 2"))
    items.append(("table", None, None))
    items.append(("heading", "Key Risks", "Heading 2"))
    items.append(("para", "The highest risks are SLM quality gap against the teacher, benchmark over-scope, protocol churn, diagram/code density affecting reviewability, and evaluation reproducibility. The updated design treats DPO, broad A2A hardening, and production-grade HA as conditional or stretch unless needed for the thesis claim.", None))
    items.append(("break", None, None))

    for kind, text, style in items:
        if kind == "break":
            p = insert_paragraph_before(anchor)
            p.add_run().add_break(WD_BREAK.PAGE)
        elif kind == "heading":
            insert_paragraph_before(anchor, text, style)
        elif kind == "para":
            insert_paragraph_before(anchor, text)
        elif kind == "table":
            p = insert_paragraph_before(anchor, "")
            body = p._p
            table = doc.add_table(rows=1, cols=4)
            try:
                table.style = "Table Grid"
            except KeyError:
                pass
            body.addnext(table._tbl)
            rows = [
                ["Area", "Measure", "MVP threshold", "Evidence artifact"],
                ["Benchmark", "MCP-TaskBench aggregate", ">= 0.70", "Pinned task suite, traces, JSONL scores"],
                ["Tool calling", "BFCL function-call F1", ">= 0.82 final gate", "Evaluation harness report"],
                ["Faithfulness", "NLI/AlignScore-style metric", ">= 0.78", "Claim-level critic logs"],
                ["Latency", "Single-hop query p50", "<= 3.0 s", "OpenTelemetry dashboard export"],
            ]
            for i, v in enumerate(rows[0]):
                set_cell_text(table.rows[0].cells[i], v, bold=True)
                shade_cell(table.rows[0].cells[i], LIGHT_BLUE)
            for row in rows[1:]:
                cells = table.add_row().cells
                for i, v in enumerate(row):
                    set_cell_text(cells[i], v)


def append_back_matter(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("Appendix D - Claim-to-Evidence Matrix", level=1)
    rows = [
        ["Research claim", "Evidence", "Baseline/comparator", "Statistical treatment"],
        ["MCP-TaskBench discriminates MCP-capable agents", "Task category scores and call traces", "BFCL, AgentBench, ToolBench where applicable", "Effect size plus bootstrapped 95% CI"],
        ["Distilled Phi-4-mini remains useful for MCP tool-calling", "BFCL F1, MCP-TaskBench aggregate, per-category failures", "Teacher model and non-distilled Phi-4-mini", "Paired tests on identical task sets"],
        ["Graph + vector retrieval improves grounded answers", "Ablation table for vector-only, graph-only, hybrid", "Vector-only RAG baseline", "Wilcoxon signed-rank and effect size"],
        ["Bounded critic loop improves faithfulness without runaway latency", "Faithfulness score, re-retrieve count, latency histogram", "No-critic and unbounded critic variants", "Paired delta plus p95 latency comparison"],
        ["Drift-triggered re-indexing preserves freshness efficiently", "Staleness lag, re-index count, stale-entity precision", "Scheduled full re-index", "Bootstrap CI over simulated commits"],
    ]
    add_table.doc = doc
    add_table(rows, widths=[2.1, 2.0, 1.8, 1.8])

    doc.add_heading("Appendix E - Operational Runbook Starters", level=1)
    rows = [
        ["Alert", "Likely cause", "First checks", "Recovery action"],
        ["query_latency_seconds p95 > 2x SLO", "vLLM saturation, pgvector HNSW regression, graph query fan-out", "Check vLLM queue, ANN latency, Neo4j slow query log", "Reduce k, disable rerank temporarily, scale inference replica"],
        ["ingest_latency p95 > 2x SLO", "Kafka backlog, parser errors, object-store fetch slowness", "Check consumer lag, parser WARN rate, MinIO latency", "Increase parser workers, replay failed partition, quarantine bad repo"],
        ["faithfulness_score p50 < 0.78", "Retrieval drift, critic threshold mismatch, stale embeddings", "Sample unsupported claims, inspect retrieval evidence", "Trigger selective re-index; retune critic threshold on dev set"],
        ["reindex_lag_seconds p95 > 5 min", "Noisy tenant monopolising slots", "Check per-tenant queue share and retry rate", "Apply tenant rate cap; move failed jobs to repair DLQ"],
    ]
    add_table(rows, widths=[1.7, 2.1, 2.0, 2.0])

    doc.add_heading("Appendix F - Key External References", level=1)
    refs = [
        ("Model Context Protocol specification", "https://modelcontextprotocol.io/specification/"),
        ("MCP releases", "https://github.com/modelcontextprotocol/modelcontextprotocol/releases"),
        ("A2A Protocol v1.0 announcement", "https://a2a-protocol.org/latest/announcing-1.0/"),
        ("Microsoft Phi-4-mini technical report", "https://www.microsoft.com/en-us/research/publication/phi-4-mini-technical-report-compact-yet-powerful-multimodal-language-models-via-mixture-of-loras/"),
        ("Neo4j documentation", "https://neo4j.com/docs/"),
        ("pgvector project", "https://github.com/pgvector/pgvector"),
        ("vLLM documentation", "https://docs.vllm.ai/"),
        ("LangGraph documentation", "https://langchain-ai.github.io/langgraph/"),
        ("Microsoft AutoGen", "https://microsoft.github.io/autogen/"),
        ("Berkeley Function Calling Leaderboard", "https://gorilla.cs.berkeley.edu/leaderboard.html"),
    ]
    for title, url in refs:
        p = doc.add_paragraph(style="List Paragraph")
        add_hyperlink(p, title, url)


def apply_text_fixes(doc: Document) -> None:
    replacements = {
        "MCP 1.0": "the current MCP specification line",
        "since MCP 1.0 was released": "as MCP matured through its date-based specification revisions",
        "a2a-sdk is stable": "the A2A v1.0 SDK line is treated as the stable target",
    }
    for p in doc.paragraphs:
        for old, new in replacements.items():
            if old in p.text:
                for run in p.runs:
                    run.text = run.text.replace(old, new)


def main() -> None:
    shutil.copy2(SRC, OUT)
    doc = Document(OUT)
    style_document(doc)
    apply_text_fixes(doc)
    append_front_matter(doc)
    append_back_matter(doc)
    doc.save(OUT)
    set_toc_update_flag(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
